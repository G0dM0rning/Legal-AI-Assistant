import os
import logging
from typing import List, Dict, Any, Optional
from langchain_core.documents import Document
from langchain.chains.summarize import load_summarize_chain
from langchain.prompts import PromptTemplate
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from app.config import llm

logger = logging.getLogger(__name__)

class LegalSummarizer:
    """Professional summarizer optimized for long Pakistani legal documents"""
    
    @property
    def llm(self):
        from app.config import ai_provider
        return ai_provider.llm

    def _is_langchain_runnable(self, llm_instance) -> bool:
        """Check if an LLM is a proper LangChain Runnable (not CustomGroqLLM)"""
        try:
            from langchain_core.runnables import Runnable
            return isinstance(llm_instance, Runnable)
        except ImportError:
            return hasattr(llm_instance, '_generate')

    def _get_prompt_text(self, summary_type: str) -> str:
        """Get the appropriate prompt template text"""
        if summary_type == "executive":
            return """
            Write a professional Executive Summary of the following legal decision/document.
            
            FORMATTING RULES:
            1. Use a clear title and structured sections.
            2. Focus on: Core Legal Issue, Key Evidence/Precedents, and the Final Ruling.
            3. DO NOT use citation markers like [Ref 1], [Ref 2], etc.
            4. Keep it clear for a layman while maintaining high legal accuracy.
            
            TEXT: {text}
            
            EXECUTIVE SUMMARY:"""
        else:
            return """
            Provide a Detailed Legal Analysis of the following document.
            
            STRICT RULES:
            - Address: Parties, Procedural History, Key Arguments, Legal Reasoning (Ratio Decidendi), and Final Order.
            - DO NOT include RAG citation markers such as [Ref X].
            - Maintain a sophisticated, expert legal tone.
            
            TEXT: {text}
            
            DETAILED ANALYSIS:"""

    def summarize_long_document(self, docs: List[Document], summary_type: str = "executive") -> str:
        """Summarize long legal documents using Map-Reduce chain or direct LLM call"""
        if not self.llm:
            return "Summarization service unavailable."
        
        combine_template = self._get_prompt_text(summary_type)
        map_template = "Summarize the following section of a legal document, preserving all key legal citations and parties mentioned: {text}"
        
        # Strategy 1: Use LangChain chain if LLM is a proper Runnable
        if self._is_langchain_runnable(self.llm):
            try:
                map_prompt = PromptTemplate(template=map_template, input_variables=["text"])
                combine_prompt = PromptTemplate(template=combine_template, input_variables=["text"])
                chain = load_summarize_chain(
                    self.llm, 
                    chain_type="map_reduce",
                    map_prompt=map_prompt,
                    combine_prompt=combine_prompt,
                    verbose=False
                )
                summary = chain.invoke(docs)
                # chain.invoke returns a dict with 'output_text' key
                if isinstance(summary, dict):
                    return summary.get('output_text', str(summary))
                return str(summary)
            except Exception as e:
                logger.error(f"Chain summarization failed: {e}")
                # Fall through to direct strategy
        
        # Strategy 2: Direct LLM call (for CustomGroqLLM or chain failure)
        try:
            logger.info("Using direct LLM call for summarization (non-Runnable LLM)")
            
            # Map phase: summarize each chunk individually
            chunk_summaries = []
            for i, doc in enumerate(docs):
                map_input = map_template.replace("{text}", doc.page_content[:4000])
                result = self.llm.invoke(map_input)
                content = result.content if hasattr(result, 'content') else str(result)
                chunk_summaries.append(content)
                logger.debug(f"Mapped chunk {i+1}/{len(docs)}")
            
            # Combine phase: merge all summaries
            combined_text = "\n\n".join(chunk_summaries)
            # Cap to avoid token limits
            if len(combined_text) > 15000:
                combined_text = combined_text[:15000]
            
            final_prompt = combine_template.replace("{text}", combined_text)
            result = self.llm.invoke(final_prompt)
            content = result.content if hasattr(result, 'content') else str(result)
            return content
        except Exception as e:
            logger.error(f"Direct summarization failed: {e}")
            return f"Failed to generate summary: {str(e)}"

    def export_to_docx(self, content: str, title: str, output_path: str):
        """Export summary to a professional Word document"""
        doc = DocxDocument()
        doc.add_heading(title, 0)
        
        doc.add_paragraph("Legal AI - Professional Case Summary")
        doc.add_paragraph(f"Document: {title}")
        doc.add_paragraph("-" * 20)
        
        doc.add_paragraph(content)
        
        doc.save(output_path)
        return output_path

    def export_to_pdf(self, content: str, title: str, output_path: str):
        """Export summary to a professional PDF document"""
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Heading
        story.append(Paragraph(f"<b>{title}</b>", styles['Heading1']))
        story.append(Paragraph("<i>Legal AI Professional Case Summary</i>", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Content
        # Split by newlines and add paragraphs
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))
        
        doc.build(story)
        return output_path

summarizer = LegalSummarizer()
